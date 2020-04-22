import docx
#run pip install python-docx
#documentation: https://python-docx.readthedocs.org

d = docx.Document('D:\Documents\Hobby\Automate the Boring Stuff With Python\Word_Files_To_Edit\demo.docx')
d.paragraphs  #A [aragraph comes after every Enter.
#[<docx.text.paragraph.Paragraph object at 0x0000023E9C3DBF70>, <docx.text.paragraph.Paragraph object at 0x0000023E9C3F4460>, <docx.text.paragraph.Paragraph object at 0x0000023E9C3F40D0>, <docx.text.paragraph.Paragraph object at 0x0000023E9C3F4220>, <docx.text.paragraph.Paragraph object at 0x0000023E9C3F4160>, <docx.text.paragraph.Paragraph object at 0x0000023E9C3F41C0>, <docx.text.paragraph.Paragraph object at 0x0000023E9C3F4040>]
d.paragraphs[0]
#<docx.text.paragraph.Paragraph object at 0x0000023E9C3F4280>
d.paragraphs[0].text
#'Document Title'
d.paragraphs[1].text
#'A plain paragraph having some bold and some italic.'

p = d.paragraphs[1]
p.runs     #A run comes after every text format change.
#[<docx.text.run.Run object at 0x0000023E9C3DBD90>, <docx.text.run.Run object at 0x0000023E9C3F4130>, <docx.text.run.Run object at 0x0000023E9C3F4190>, <docx.text.run.Run object at 0x0000023E9C3F4100>]
p.runs[0].text
#'A plain paragraph having some '
p.runs[1].text
#'bold'
p.runs[2].text
#' and some '
p.runs[3].text
#'italic.'
p.runs[1].bold
#True
p.runs[0].bold
p.runs[0].bold == None
#True

p.runs[3].italic
#True
p.runs[3].underline = True    #You can edit the file like this.
p.runs[3].text = 'italic and underlined.'
d.save('D:\Documents\Hobby\Automate the Boring Stuff With Python\Word_Files_To_Edit\demo2.docx')
#remember, any changes you make is all on memory. If you want to sae it, you need to run this.

p.style
#_ParagraphStyle('Normal') id: 2467932619920
p.style = 'Title'
p.style
#_ParagraphStyle('Title') id: 2467941216896
d.save('D:\Documents\Hobby\Automate the Boring Stuff With Python\Word_Files_To_Edit\demo3.docx')



d = docx.Document()   #start a new document, still just in memory.
d.add_paragraph('Hello, this is a paragraph.')
#<docx.text.paragraph.Paragraph object at 0x0000023E9CC27580>
d.add_paragraph('This is another paragraph.')
#<docx.text.paragraph.Paragraph object at 0x0000023E9C3DBD90>
d.save('D:\Documents\Hobby\Automate the Boring Stuff With Python\Word_Files_To_Edit\demo4.docx')

p=d.paragraphs[0]
p.add_run('This is a new run.')
#<docx.text.run.Run object at 0x0000023E9CC272E0>
p.runs
#[<docx.text.run.Run object at 0x0000023E9CC272B0>, <docx.text.run.Run object at 0x0000023E9CC273A0>]
p.runs[1].bold = True
d.save('D:\Documents\Hobby\Automate the Boring Stuff With Python\Word_Files_To_Edit\demo5.docx')

